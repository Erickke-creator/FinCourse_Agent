"""Generate 项目使用与协作手册.docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ---- 全局样式设置 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
# Set East Asian font
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), '微软雅黑')
rPr.insert(0, rFonts)

# Configure heading styles
for i, (size, color_hex) in enumerate([(22, '1a1a2e'), (16, '16213e'), (14, '0f3460'), (12, '533483')], 1):
    h_style = doc.styles[f'Heading {i}']
    h_style.font.size = Pt(size)
    h_style.font.bold = True
    h_style.font.color.rgb = RGBColor(*tuple(int(color_hex[j:j+2], 16) for j in (0, 2, 4)))
    h_style.font.name = '微软雅黑'
    h_style.paragraph_format.space_before = Pt(16 if i > 1 else 24)
    h_style.paragraph_format.space_after = Pt(8)

# ---- Helper functions ----
def add_code_block(doc, text):
    """Add a code block in Consolas font with grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    # Add shading
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0',
    })
    pPr.insert(0, shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    return p

def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacer
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_note(doc, text):
    """Add a highlighted note."""
    p = doc.add_paragraph()
    run = p.add_run('⚠ ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    run.font.italic = True
    return p

# ================================================================
# 封面
# ================================================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('小微企业贷款智能评估 Agent')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('项目使用与协作手册')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x53, 0x34, 0x83)

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.add_run('金融科技课程项目 | 2026年6月').font.size = Pt(12)

doc.add_page_break()

# ================================================================
# 目录页（手动提示）
# ================================================================
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述',
    '二、环境准备与启动',
    '    2.1 克隆仓库',
    '    2.2 安装依赖',
    '    2.3 训练 ML 模型',
    '    2.4 启动项目',
    '三、Git 基础操作',
    '    3.1 拉取最新代码',
    '    3.2 提交修改',
    '    3.3 推送到远程仓库',
    '    3.4 冲突处理',
    '四、仓库管理工作流',
    '    4.1 分支策略',
    '    4.2 Commit 信息规范',
    '    4.3 知识库更新流程',
    '    4.4 Pull Request 流程',
    '五、知识库协作规范',
    '六、常见问题与解决',
    '七、附录：项目目录结构',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ================================================================
# 一、项目概述
# ================================================================
doc.add_heading('一、项目概述', level=1)

doc.add_paragraph(
    '本项目是一个面向小微企业主的贷款智能评估系统。用户输入企业经营数据后，'
    '系统通过规则引擎 + ML 模型 + 知识库检索，输出信用评分、风险诊断、'
    '26 家银行的贷款通过概率预测和改善建议。'
)

doc.add_paragraph('技术栈：')
add_bullet(doc, '后端：Python + FastAPI + XGBoost / GradientBoosting / RandomForest')
add_bullet(doc, '前端：React 19 + Vite + TypeScript + TailwindCSS + ECharts')
add_bullet(doc, '知识库：结构化 CSV/JSON（16 个数据文件）+ Python Loader 包')
add_bullet(doc, '协作：Git + GitHub')

doc.add_heading('系统架构', level=2)
doc.add_paragraph(
    '用户通过前端页面输入企业数据 → 前端调用后端 /api/evaluate 接口 → '
    '后端从统一知识库（kb/data/）加载银行产品、行业准入、政策规则等数据 → '
    '规则引擎计算五维评分 + 银行匹配 → ML 模型预测违约概率和信用评级 → '
    '返回评估结果 + 知识库溯源信息 → 前端展示评分仪表盘、银行排名、改善建议。'
)

doc.add_page_break()

# ================================================================
# 二、环境准备与启动
# ================================================================
doc.add_heading('二、环境准备与启动', level=1)

doc.add_heading('2.1 克隆仓库', level=2)
doc.add_paragraph('在你要存放项目的文件夹中打开终端（Git Bash），执行：')
add_code_block(doc, 'git clone https://github.com/Erickke-creator/FinCourse_Agent.git')
add_code_block(doc, 'cd FinCourse_Agent')
doc.add_paragraph('如果提示"command not found"，请先安装 Git：https://git-scm.com/download/win')

doc.add_heading('2.2 安装依赖', level=2)
doc.add_paragraph('Python 后端（需要 Python ≥ 3.9）：')
add_code_block(doc, 'cd "找资料/FinTech_小微企业贷款评估系统/FinTech_小微企业贷款评估系统/后端服务"')
add_code_block(doc, 'pip install -r requirements.txt')
doc.add_paragraph('')  # spacer
doc.add_paragraph('Node 前端（需要 Node.js ≥ 18）：')
add_code_block(doc, 'cd "../前端源码"')
add_code_block(doc, 'npm install')

doc.add_heading('2.3 训练 ML 模型', level=2)
doc.add_paragraph('首次使用必须执行。在后端服务目录下运行：')
add_code_block(doc, 'cd "../后端服务"')
add_code_block(doc, 'python train_ml_enhanced.py')
doc.add_paragraph('训练完成后会在 models/ 目录下生成 .pkl 模型文件（已在 .gitignore 中排除）。')

doc.add_heading('2.4 启动项目', level=2)
doc.add_paragraph('第一步：启动后端（端口 8000）')
add_code_block(doc, '# 终端 1')
add_code_block(doc, 'cd "后端服务"')
add_code_block(doc, 'python -m uvicorn main:app --host 0.0.0.0 --port 8000')
doc.add_paragraph('看到 "Knowledge Base loaded" 和 "Uvicorn running on http://0.0.0.0:8000" 表示启动成功。')
doc.add_paragraph('')  # spacer
doc.add_paragraph('第二步：启动前端（端口 3000）')
add_code_block(doc, '# 终端 2')
add_code_block(doc, 'cd "前端源码"')
add_code_block(doc, 'npm run dev')
doc.add_paragraph('浏览器打开 http://localhost:3000，选择一个案例或手动填写企业数据，点击"开始智能评估 + 银行匹配"，确认能看到评分结果和银行排名。')

doc.add_page_break()

# ================================================================
# 三、Git 基础操作
# ================================================================
doc.add_heading('三、Git 基础操作', level=1)

doc.add_paragraph('以下操作均在项目根目录（FinCourse_Agent）下执行。')

doc.add_heading('3.1 拉取最新代码（每次工作前）', level=2)
add_code_block(doc, 'git pull')
doc.add_paragraph('这条命令会把远程仓库上队友的最新修改同步到你的本地。建议每次打开项目都先 pull 一下。')

doc.add_heading('3.2 查看当前状态', level=2)
add_code_block(doc, 'git status')
doc.add_paragraph('红色文件 = 已修改但未暂存 | 绿色文件 = 已暂存待提交')

doc.add_heading('3.3 提交修改（三步走）', level=2)
doc.add_paragraph('第一步：将修改的文件加入暂存区')
add_code_block(doc, 'git add <文件名>      # 添加指定文件')
add_code_block(doc, 'git add .              # 添加当前目录下所有修改')
doc.add_paragraph('第二步：创建 commit')
add_code_block(doc, 'git commit -m "类型: 改动说明"')
doc.add_paragraph('第三步：推送到远程仓库')
add_code_block(doc, 'git push')
add_note(doc, '如果 push 被拒绝（别人先推送了），先执行 git pull --rebase，再 git push。')

doc.add_heading('3.4 修改知识库的完整示例', level=2)
doc.add_paragraph('场景：修改制造业的银行接受度系数从 1.0 改为 1.05')
add_code_block(doc, '# 1. 编辑文件')
add_code_block(doc, '# 打开 找资料/kb/data/industries/industry_acceptance.csv')
add_code_block(doc, '# 把 manufacturing 那行的接受度系数改为 1.05')
add_code_block(doc, '')
add_code_block(doc, '# 2. 更新版本日期')
add_code_block(doc, '# 打开 找资料/kb/VERSION，把 date 改成当天')
add_code_block(doc, '')
add_code_block(doc, '# 3. 提交')
add_code_block(doc, 'git add "找资料/kb/data/industries/industry_acceptance.csv" "找资料/kb/VERSION"')
add_code_block(doc, 'git commit -m "kb: 制造业接受度系数 1.0→1.05，依据2026央行指导"')
add_code_block(doc, 'git push')

doc.add_heading('3.5 冲突处理', level=2)
doc.add_paragraph('当两个人都修改了同一文件时，Git 无法自动合并，需要手动解决：')
add_code_block(doc, '# 拉取并变基')
add_code_block(doc, 'git pull --rebase')
add_code_block(doc, '')
add_code_block(doc, '# 如果有冲突，会提示 CONFLICT')
add_code_block(doc, '# 打开冲突文件，找到 <<<<<<< ======= >>>>>>> 标记')
add_code_block(doc, '# 手动选择保留哪些内容，删除标记线')
add_code_block(doc, '')
add_code_block(doc, '# 解决后继续')
add_code_block(doc, 'git add .')
add_code_block(doc, 'git rebase --continue')
add_code_block(doc, 'git push')

doc.add_page_break()

# ================================================================
# 四、仓库管理工作流
# ================================================================
doc.add_heading('四、仓库管理工作流', level=1)

doc.add_heading('4.1 分支策略', level=2)
doc.add_paragraph(
    '我们使用简化版的分支策略，适合学生团队：'
)
add_table(doc,
    ['分支', '用途', '谁可以直接推送'],
    [
        ['main', '稳定版本（可演示，可交作业）', '全体（需 review）'],
        ['<人名>/<功能>', '个人开发分支（如 zhangsan/fix-bank-engine）', '该成员本人'],
    ]
)

doc.add_paragraph('工作流程：')
add_bullet(doc, '从 main 创建你的个人分支：git checkout -b zhangsan/update-policy')
add_bullet(doc, '在你的分支上开发、提交、推送')
add_bullet(doc, '在 GitHub 上发起 Pull Request，请组长 review')
add_bullet(doc, '组长确认没问题后合并到 main')
add_bullet(doc, '删除你的个人分支：git branch -d zhangsan/update-policy')

doc.add_heading('4.2 Commit 信息规范', level=2)
add_table(doc,
    ['前缀', '用途', '示例'],
    [
        ['kb:', '知识库数据更新', 'kb: 广东省制造业准入系数调整'],
        ['feat:', '新功能', 'feat: 增加供应链关系图谱'],
        ['fix:', 'Bug 修复', 'fix: 修正 D 级纳税评分错误'],
        ['refactor:', '代码重构', 'refactor: 抽取银行匹配引擎'],
        ['docs:', '文档更新', 'docs: 更新 README'],
        ['chore:', '杂项/清理', 'chore: 清理旧版知识库文件'],
    ]
)

doc.add_heading('4.3 知识库更新流程', level=2)
doc.add_paragraph('知识库是项目的核心数据资产，更新时务必遵守以下流程：')
add_bullet(doc, '1. 确定要修改的数据文件（都在 找资料/kb/data/ 下）')
add_bullet(doc, '2. 用任意文本编辑器或 Excel 打开 CSV，表格编辑器打开 JSON')
add_bullet(doc, '3. 修改数据，确保格式正确（CSV 用 UTF-8 编码，JSON 用 UTF-8）')
add_bullet(doc, '4. 更新 找资料/kb/VERSION 中的 date 字段为当天日期')
add_bullet(doc, '5. 运行 python 找资料/kb/loader/loader.py 确认加载无报错')
add_bullet(doc, '6. git add + git commit -m "kb: ..." + git push')
add_note(doc, '严禁修改旧知识库目录（Agent整合构建/、普惠金融Agent_知识库/、金融科技agent/），这些是归档副本。所有新修改只应在 kb/data/ 下进行。')

doc.add_heading('4.4 Pull Request（PR）流程', level=2)
doc.add_paragraph('在 GitHub 网页上操作：')
add_bullet(doc, '1. 推送你的分支到 GitHub：git push origin <你的分支名>')
add_bullet(doc, '2. 打开 GitHub 仓库页面，点击 "Compare & pull request"')
add_bullet(doc, '3. 写清楚改了什么、为什么改')
add_bullet(doc, '4. 在右侧 Assignees 指定 reviewer（通常是组长）')
add_bullet(doc, '5. 组长 review 后点击 "Merge pull request"')

doc.add_page_break()

# ================================================================
# 五、知识库协作规范
# ================================================================
doc.add_heading('五、知识库协作规范', level=1)

doc.add_paragraph('知识库目录结构（找资料/kb/data/）：')
add_table(doc,
    ['目录', '内容', '文件数', '格式'],
    [
        ['policies/', '国家级 + 地方级政策规则', '2', 'CSV'],
        ['banks/', '28家银行产品 + 地域可用性', '2', 'JSON + CSV'],
        ['industries/', '18行业准入 + 地域调整系数', '2', 'CSV'],
        ['credit_tax/', '征信容忍度 + 纳税评分', '2', 'CSV'],
        ['risk_control/', '被拒因子 + 补贴政策 + 宏观数据', '3', 'CSV + JSON'],
        ['cases/', '30条基础案例 + 20条增强案例', '2', 'CSV'],
        ['governance/', '数据来源 + 语义说明 + 字段映射', '3', 'CSV + MD'],
    ]
)

doc.add_heading('修改规则', level=2)
add_bullet(doc, '只能修改 kb/data/ 下的文件，不能修改旧知识库目录')
add_bullet(doc, 'CSV 文件必须保持 UTF-8 编码，列名不能改')
add_bullet(doc, 'JSON 文件必须保持合法 JSON 格式（可以用 jsonlint.com 校验）')
add_bullet(doc, '每次修改后必须更新 kb/VERSION 中的日期')
add_bullet(doc, '修改后重启后端服务，使新数据生效')

doc.add_page_break()

# ================================================================
# 六、常见问题
# ================================================================
doc.add_heading('六、常见问题与解决', level=1)

doc.add_heading('Q1: push 时提示 "fatal: unable to access"', level=2)
doc.add_paragraph('网络问题或 GitHub 认证过期。解决：')
add_bullet(doc, '检查网络连接，尝试 ping github.com')
add_bullet(doc, '如果是认证问题：git config --global credential.helper manager')

doc.add_heading('Q2: commit 时提示 "Please tell me who you are"', level=2)
doc.add_paragraph('首次使用需要设置身份：')
add_code_block(doc, 'git config user.name "你的名字"')
add_code_block(doc, 'git config user.email "你的邮箱@qq.com"')

doc.add_heading('Q3: 误删了文件想恢复', level=2)
add_code_block(doc, 'git restore <文件路径>      # 恢复到最近一次 commit 的状态')

doc.add_heading('Q4: 不小心 commit 了不该提交的文件', level=2)
add_code_block(doc, '# 从 Git 追踪中移除（文件保留在本地）')
add_code_block(doc, 'git rm --cached <文件路径>')
add_code_block(doc, 'git commit -m "chore: 移除误提交的文件"')
add_code_block(doc, '')
add_code_block(doc, '# 同时别忘记更新 .gitignore 防止再次误提交')

doc.add_heading('Q5: 想撤回最近的 commit（但还没 push）', level=2)
add_code_block(doc, 'git reset --soft HEAD~1    # 撤销 commit，保留文件修改')
add_code_block(doc, 'git reset --hard HEAD~1    # 撤销 commit，丢弃所有修改（危险！）')

doc.add_heading('Q6: 后端启动后报 "Knowledge Base not available"', level=2)
doc.add_paragraph('这是正常的退路提示，表示 kb/loader 包未被正确识别。评估仍会使用硬编码的 fallback 数据完成，不影响使用。如需修复：')
add_code_block(doc, 'cd 找资料/kb/loader')
add_code_block(doc, 'pip install -e .')

doc.add_heading('Q7: 前端看不到"知识库溯源"面板', level=2)
doc.add_paragraph('说明后端不在线或使用了 fallback 模式。确认后端已启动且终端日志显示"Knowledge Base loaded"后，刷新前端页面。')

doc.add_heading('Q8: 想回退到某个历史版本查看之前的代码', level=2)
add_code_block(doc, 'git log --oneline             # 查看 commit 历史')
add_code_block(doc, 'git checkout <commit-hash>    # 临时切换到该版本（只读）')
add_code_block(doc, 'git checkout main             # 回到最新版本')

doc.add_page_break()

# ================================================================
# 七、附录
# ================================================================
doc.add_heading('七、附录：项目目录结构', level=1)

add_code_block(doc, '''FinCourse_Agent/
├── .gitignore              # Git 忽略规则
├── GIT_GUIDE.md            # Git 协作手册
├── 找资料/
│   ├── kb/                 # ★ 统一知识库（唯一权威来源）
│   │   ├── VERSION         #   版本号
│   │   ├── README.md       #   使用说明
│   │   ├── data/           #   数据文件（CSV + JSON）
│   │   │   ├── policies/   #   政策规则
│   │   │   ├── banks/      #   银行产品
│   │   │   ├── industries/ #   行业准入
│   │   │   ├── credit_tax/ #   征信与纳税
│   │   │   ├── risk_control/#  风控规则
│   │   │   ├── cases/      #   教学案例
│   │   │   └── governance/ #   数据治理
│   │   └── loader/         #   Python 加载器包
│   ├── FinTech_小微企业贷款评估系统/  # ★ 应用代码
│   │   └── .../
│   │       ├── 后端服务/   # FastAPI + ML
│   │       ├── 前端源码/   # React + Vite
│   │       └── 数据与报告/ # JSON 参考数据
│   ├── Agent整合构建/      # 归档（只读）
│   ├── 普惠金融Agent_知识库/ # 归档（只读）
│   └── 金融科技agent/      # 归档（只读）''')

doc.add_paragraph()
doc.add_paragraph('—— 文档结束 ——').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---- Save ----
output_path = os.path.join(
    os.path.dirname(__file__), '..', '..',
    '项目使用与协作手册.docx'
)
output_path = os.path.abspath(output_path)
doc.save(output_path)
print(f'Saved: {output_path}')
