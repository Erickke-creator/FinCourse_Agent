"""
生成 FinTech 小微贷款智能评估系统 v5 — 项目使用与协作手册 (Word)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "项目协作手册_v5.docx")

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("FinTech 小微贷款智能评估系统 v5")
run.font.size = Pt(24)
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("项目使用与协作手册").font.size = Pt(16)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("版本: v5.0 | 日期: 2026-06-26 | 基于 FastAPI + React + DeepSeek V4").font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 1. 项目概述
# ============================================================
doc.add_heading("一、项目概述", level=1)
doc.add_paragraph(
    "FinTech 小微贷款智能评估系统是一个面向小微企业主的贷款可行性自评工具。"
    "v5 版本核心升级：将原来基于关键词匹配的对话机器人替换为 DeepSeek V4 LLM Agent，"
    "支持 Function Call 实时查询知识库、评估贷款风险、匹配银行产品，"
    "并在评估流程中集成了 XGBoost/GB/RF 等 ML 模型预测结果。"
)

# ============================================================
# 2. 目录结构
# ============================================================
doc.add_heading("二、目录结构", level=1)

doc.add_paragraph(
    "仓库根目录下，所有版本代码集中在「各版本代码」文件夹中："
)

tree = """各版本代码/
├── kb/                            # ★ 统一知识库（所有版本共用）
│   ├── data/                      #   16个数据文件（CSV/JSON/Markdown）
│   │   ├── policies/              #   国家级 + 地方级政策规则
│   │   ├── banks/                 #   银行产品数据库（28家）
│   │   ├── industries/            #   行业准入规则
│   │   ├── credit_tax/            #   征信容忍度 + 纳税评分
│   │   ├── risk_control/          #   被拒因子 + 补贴政策 + 宏观数据
│   │   ├── cases/                 #   教学案例（基础30 + 增强20）
│   │   └── governance/            #   数据治理（来源登记 + 语义说明）
│   ├── loader/                    #   Python 知识库加载器（pip 包）
│   └── tools/                     #   文档/数据迁移工具脚本
│
├── FinTech_小微贷款智能评估系统_v4/  # v4 版本（规则引擎，保留参考）
│   └── FinTech_小微贷款智能评估系统_v4/
│       ├── 后端服务/               #   FastAPI 后端
│       ├── 前端源码/               #   React + Vite 前端
│       └── 数据文件/               #   研究与参考报告
│
├── FinTech_小微贷款智能评估系统_v5/  # ★ v5 当前开发版本
│   ├── .env.example               #   API Key 配置模版
│   ├── 后端服务/                   #   FastAPI 后端
│   │   ├── main.py                #     API 入口
│   │   ├── chat_agent.py          # ★   LLM Agent 循环（DeepSeek V4）
│   │   ├── kb_bridge.py           # ★   知识库桥接层（新增）
│   │   ├── bank_engine.py         # ★   评估引擎（已集成 ML）
│   │   ├── models.py              #     数据模型
│   │   ├── ml_inference.py        #     ML 推理封装
│   │   ├── enterprise_search.py   #     企业搜索
│   │   ├── cumcm_data/            #     训练数据
│   │   ├── models/                #     ML 模型文件（.pkl，不提交 Git）
│   │   └── requirements.txt       #     Python 依赖
│   ├── 前端源码/                   #   React + Vite + TypeScript 前端
│   └── 数据文件/                   #   研究参考报告
│
├── Agent功能改进路线图.docx          # 功能迭代规划文档
└── 项目协作手册_v5.docx             # 本手册"""

doc.add_paragraph(tree)

# ============================================================
# 3. 环境准备
# ============================================================
doc.add_heading("三、环境准备", level=1)

doc.add_heading("3.1 安装 Python 依赖", level=2)
doc.add_paragraph(
    "进入 v5 后端目录，安装依赖：\n"
    "cd 各版本代码\\FinTech_小微贷款智能评估系统_v5\\后端服务\n"
    "pip install -r requirements.txt"
)

doc.add_heading("3.2 训练 ML 模型（首次必做）", level=2)
doc.add_paragraph(
    "在同目录下运行训练脚本，生成 .pkl 模型文件：\n"
    "python train_ml_enhanced.py\n\n"
    "该脚本会输出 XGBoost、GradientBoosting、RandomForest 三个模型到 models/ 目录。\n"
    "⚠️ .pkl 文件不会被提交到 Git（已在 .gitignore 中排除），每位成员需自行训练。"
)

doc.add_heading("3.3 配置 DeepSeek API Key", level=2)
doc.add_paragraph(
    "复制 .env.example 为 .env，并填入真实 API Key：\n"
    "copy ..\\.env.example ..\\.env\n"
    "用文本编辑器打开 .env，将 sk-your-key-here 替换为真实 Key。\n\n"
    "获取方式：https://platform.deepseek.com → 注册 → API Keys → 创建\n"
    "最低充值 ¥1，按量计费（¥1/百万 token），开发测试阶段 ¥10 足够。\n"
    "⚠️ .env 文件已在 .gitignore 中排除，不会被提交到 Git。"
)

doc.add_heading("3.4 安装前端依赖", level=2)
doc.add_paragraph(
    "cd 各版本代码\\FinTech_小微贷款智能评估系统_v5\\前端源码\n"
    "npm install"
)

doc.add_heading("3.5 配置 Git 代理（中国大陆必做）", level=2)
doc.add_paragraph(
    "由于 GitHub 在中国大陆访问受限，需配置代理。以 Clash 为例（端口在客户端首页查看）：\n"
    "git config --global http.proxy http://127.0.0.1:7897\n"
    "git config --global https.proxy http://127.0.0.1:7897\n\n"
    "其他 VPN 同理，替换为对应端口即可。配置后验证：\n"
    "git config --global --get http.proxy"
)

# ============================================================
# 4. 启动运行
# ============================================================
doc.add_heading("四、启动运行", level=1)

doc.add_heading("4.1 启动后端", level=2)
doc.add_paragraph(
    "cd 各版本代码\\FinTech_小微贷款智能评估系统_v5\\后端服务\n"
    "python main.py\n\n"
    "成功标志：终端显示 \"AI Mode: Enabled (DeepSeek V4)\"。\n"
    "API 文档：http://localhost:8000/docs"
)

doc.add_heading("4.2 启动前端", level=2)
doc.add_paragraph(
    "cd 各版本代码\\FinTech_小微贷款智能评估系统_v5\\前端源码\n"
    "npm run dev\n\n"
    "浏览器打开 http://localhost:3000，即可使用。"
)

doc.add_heading("4.3 API 端点速查", level=2)

api_table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
headers = ['端点', '方法', '用途']
for i, h in enumerate(headers):
    api_table.rows[0].cells[i].text = h
data = [
    ['/api/health', 'GET', '健康检查（含 AI 状态）'],
    ['/api/evaluate', 'POST', '贷款评估（含 ML 预测）'],
    ['/api/chat', 'POST', 'AI 对话（LLM Agent + 知识库查询）'],
    ['/api/chat/reset', 'POST', '重置对话会话'],
    ['/api/banks', 'GET', '银行产品列表'],
    ['/api/enterprise/search', 'POST', '企业模糊搜索'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        api_table.rows[r+1].cells[c].text = val

# ============================================================
# 5. 知识库协作规则
# ============================================================
doc.add_heading("五、知识库协作规则", level=1)

doc.add_heading("5.1 统一知识库位置", level=2)
doc.add_paragraph(
    "所有版本共用「各版本代码/kb/data/」下的知识库文件。\n"
    "v5 后端通过 kb_bridge.py 自动读取，无需在每个版本内复制知识库副本。"
)

doc.add_heading("5.2 修改规范", level=2)
rules = [
    "只能修改 kb/data/ 下的文件，禁止在各版本目录内创建私有的知识库副本",
    "CSV 文件必须保持 UTF-8 BOM 编码（用 VS Code / Excel 另存时注意）",
    "新增政策/案例时，同步更新 kb/VERSION 文件中的版本号和文件计数",
    "修改后运行 kb_bridge.py 验证读取正常：python -c \"from kb_bridge import get_kb_summary; print(get_kb_summary())\"",
    "数据来源必须在 kb/data/governance/data_source_registry.csv 中登记",
    "Commit 前缀使用 kb:，例如：git commit -m \"kb: 新增 2026 年 XX 省贴息政策\"",
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

# ============================================================
# 6. Git 协作规范
# ============================================================
doc.add_heading("六、Git 协作规范", level=1)

doc.add_heading("6.1 Commit 前缀约定", level=2)
commit_table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['前缀', '用途', '示例']):
    commit_table.rows[0].cells[i].text = h
commit_data = [
    ['feat:', '新功能', 'feat: Chat Agent 接入 DeepSeek V4'],
    ['fix:', 'Bug 修复', 'fix: kb_bridge CSV 空值崩溃'],
    ['refactor:', '重构/优化', 'refactor: 删除废弃目录，统一知识库路径'],
    ['kb:', '知识库变更', 'kb: 更新 2026 Q2 宏观统计数据'],
    ['docs:', '文档更新', 'docs: 新增项目协作手册 v5'],
    ['chore:', '杂项/工具', 'chore: 更新 .gitignore 排除 .pkl'],
    ['style:', '格式调整', 'style: CSV 统一 UTF-8 BOM 编码'],
]
for r, row_data in enumerate(commit_data):
    for c, val in enumerate(row_data):
        commit_table.rows[r+1].cells[c].text = val

doc.add_heading("6.2 分支策略", level=2)
doc.add_paragraph(
    "main 分支：稳定版本，通过 PR 合并。\n"
    "开发直接在 main 上进行（小组规模小，无需额外分支）。\n"
    "如需试验大改，可开 feature/xxx 分支。"
)

doc.add_heading("6.3 禁止提交的内容", level=2)
banned = [
    ".env 文件（含 API Key）— 已被 .gitignore 拦截",
    "models/*.pkl（ML 模型文件）— 组员各自训练",
    "node_modules/（前端依赖）— 通过 npm install 安装",
    "__pycache__/、*.pyc（Python 缓存）",
    ".DS_Store（Mac 系统文件）",
]
for b in banned:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading("6.4 日常协作流程", level=2)
flow = [
    "git pull origin main — 开工前拉取最新代码",
    "修改代码 / 更新知识库",
    "git add -A ; git commit -m \"前缀: 说明\" — 提交",
    "git pull origin main --rebase — 推送前再拉一次，避免冲突",
    "git push origin main — 推送",
    "遇到冲突：先沟通，手动合并冲突标记，再 commit + push",
]
for i, f in enumerate(flow, 1):
    doc.add_paragraph(f"{i}. {f}")

# ============================================================
# 7. 常见问题
# ============================================================
doc.add_heading("七、常见问题", level=1)

faqs = [
    ("Q: 启动后端报错 \"API Key 不可用\"",
     "A: 正常——没有 .env 文件时 AI 对话降级为基础模式（仍可使用评估仪表盘）。创建 .env 填入 Key 后重启即可。"),
    ("Q: pip install 报错 / 找不到包",
     "A: 确认在 v5/后端服务/ 目录下运行。如遇网络问题，配 pip 镜像：pip install -r requirements.txt -i https://pypi.tuna.tsinghua.edu.cn/simple"),
    ("Q: ML 模型训练报错 \"No module named xgboost\"",
     "A: pip install xgboost 即可。训练脚本需要 pandas、numpy、scikit-learn、xgboost、joblib 五个包。"),
    ("Q: 前端 npm install 很慢",
     "A: 配淘宝镜像：npm config set registry https://registry.npmmirror.com"),
    ("Q: Chat 回复 \"AI 服务暂时不可用\"",
     "A: 检查：(1) .env 中 Key 是否正确；(2) Key 是否还有余额；(3) 代理/VPN 是否正常。"),
    ("Q: Git push 报 \"fetch first\" 被拒绝",
     "A: 先 git pull origin main --rebase 拉取远程更新，解决冲突后再 push。"),
    ("Q: git push 连不上 GitHub",
     "A: 配代理（见 3.5 节）。确认代理端口正确，打开 Clash 客户端查看。"),
    ("Q: v4 和 v5 有什么区别？",
     "A: v4 对话是纯关键词匹配（13个硬编码话题）。v5 改为 LLM Agent（DeepSeek V4 Function Call），能自动查知识库、调评估引擎、搜企业库后动态生成回复。v5 还集成了 ML 模型预测。"),
]
for q, a in faqs:
    doc.add_heading(q, level=3)
    doc.add_paragraph(a)

# ============================================================
# 附：v5 架构图
# ============================================================
doc.add_heading("八、v5 系统架构", level=1)
doc.add_paragraph(
    "┌─────────────────────────────────────────────────┐\n"
    "│  React Frontend (localhost:3000)                 │\n"
    "│  ChatPanel ─→ /api/chat (LLM Agent)              │\n"
    "│  Dashboard ─→ /api/evaluate (5-dim + ML)         │\n"
    "└──────────────────┬──────────────────────────────┘\n"
    "                   │\n"
    "┌──────────────────▼──────────────────────────────┐\n"
    "│  FastAPI Backend (localhost:8000)                │\n"
    "│                                                  │\n"
    "│  chat_agent.py  ─── LLM Agent Loop               │\n"
    "│    ├── Tool: search_kb     (知识库查询)           │\n"
    "│    ├── Tool: evaluate_loan (贷款评估 + ML)        │\n"
    "│    ├── Tool: search_enterprise (企业搜索)         │\n"
    "│    └── Tool: search_cases  (案例匹配)             │\n"
    "│                                                  │\n"
    "│  kb_bridge.py  ─── 统一知识库桥接层               │\n"
    "│  bank_engine.py ── 评估引擎 (+ ML inference)      │\n"
    "└──────────────────┬──────────────────────────────┘\n"
    "                   │\n"
    "┌──────────────────▼──────────────────────────────┐\n"
    "│  Unified KB (各版本代码/kb/data/)                 │\n"
    "│  42 policies | 28 banks | 50 cases | 18 industries│\n"
    "└─────────────────────────────────────────────────┘"
)

# ============================================================
# 保存
# ============================================================
doc.save(OUTPUT)
print(f"Document saved to: {OUTPUT}")
