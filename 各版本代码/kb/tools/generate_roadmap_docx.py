"""Generate Agent功能改进路线图.docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ---- 全局样式 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), '微软雅黑')
rPr.insert(0, rFonts)

for i, (size, hex_color) in enumerate([(22, '1a1a2e'), (16, '16213e'), (14, '0f3460'), (12, '533483')], 1):
    hs = doc.styles[f'Heading {i}']
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(*tuple(int(hex_color[j:j+2], 16) for j in (0, 2, 4)))
    hs.font.name = '微软雅黑'
    hs.paragraph_format.space_before = Pt(16 if i > 1 else 24)
    hs.paragraph_format.space_after = Pt(8)

# ---- 辅助函数 ----
def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F0F0'})
    pPr.insert(0, shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    return p

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for pp in c.paragraphs:
            for rn in pp.runs: rn.font.bold = True; rn.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for pp in c.paragraphs:
                for rn in pp.runs: rn.font.size = Pt(10)
    doc.add_paragraph()
    return t

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0: p.paragraph_format.left_indent = Cm(1.5*(level+1))
    return p

def note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
    run.font.italic = True

def warn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('⚠ ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

def star_rating(doc, n):
    """Add a visual star rating"""
    p = doc.add_paragraph()
    run = p.add_run('★' * n + '☆' * (5 - n))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xf3, 0x9c, 0x12)
    return p

# ================================================================
# 封面
# ================================================================
for _ in range(7):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('小微企业贷款智能评估 Agent')
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('功能改进路线图')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x53, 0x34, 0x83)

doc.add_paragraph()
mp = doc.add_paragraph()
mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mp.add_run('金融科技课程项目 | 2026年6月').font.size = Pt(12)

doc.add_page_break()

# ================================================================
# 总览
# ================================================================
doc.add_heading('改进方向总览', level=1)

doc.add_paragraph(
    '当前系统本质上是一个"规则引擎 + 传统 ML + 静态知识库"，'
    '虽然命名为 Agent，但缺少现代 AI Agent 的核心特征：自主推理、工具调用、自然语言交互、多轮对话。'
    '以下从五个维度提出功能改进建议，按可行性和演示效果排序。'
)

tbl(doc,
    ['优先级', '功能', '工作量', '演示效果', '核心价值'],
    [
        ['P0', 'LLM 自然语言诊断报告', '1天', '★★★★★', '从固定模板变为个性化分析'],
        ['P0', '多轮对话与 What-if 分析', '1-2天', '★★★★★', '从单次计算变为交互式顾问'],
        ['P1', 'RAG 政策检索与引用', '1天', '★★★★☆', '让每条建议都有政策依据'],
        ['P1', '语义案例匹配', '0.5天', '★★★☆☆', '用历史案例增强说服力'],
        ['P2', '评估报告 PDF 导出', '1天', '★★★☆☆', '实用价值：拿着去银行'],
        ['P2', '现金流压力测试', '0.5天', '★★★☆☆', '风险预警，更贴近真实场景'],
        ['P3', '知识库过期自动提醒', '0.5天', '★★☆☆☆', '数据治理，保持知识库时效'],
        ['Bonus', '多 Agent 协作评估', '2-3天', '★★★★☆', '课程加分项：AI Agent 方向'],
    ]
)

doc.add_page_break()

# ================================================================
# P0-1: LLM 诊断报告
# ================================================================
doc.add_heading('P0-1：接入 LLM 做自然语言诊断报告', level=1)

doc.add_heading('现状问题', level=2)
doc.add_paragraph(
    'AdvisorReport.tsx 组件中的诊断文字是 hardcoded 模板。'
    '不管什么企业，只要分数在 55-80 之间，显示的都是同一段"资质合规性完备，但流动性能动性一般"。'
    '这导致两个问题：'
)
bullet(doc, '不同行业、不同风险特征的企业拿到完全一样的话术')
bullet(doc, '无法引用具体的政策条文、案例数据来支撑诊断结论')

doc.add_heading('改进方案', level=2)
doc.add_paragraph('在后端新增 /api/evaluate/llm 端点，将评估流水线改为：')

bullet(doc, '1. 前端提交用户数据 → 规则引擎计算五维评分 + 银行匹配')
bullet(doc, '2. KBQuery 检索相关知识库条目（行业准入、政策规则、相似案例、纳税评分）')
bullet(doc, '3. 将所有结构化数据打包成一个 prompt，调用 LLM API')
bullet(doc, '4. LLM 生成个性化诊断报告，前端展示')

doc.add_heading('Prompt 设计示例', level=2)
code(doc, '''你是一位资深的小微企业融资顾问。请根据以下数据给出个性化的融资诊断报告。

## 企业画像
- 行业：制造业 | 地区：广东省 | 经营年限：5年 | 纳税等级：A
- 月营收：50万元 | 月成本：30万元 | 现有负债：10万元
- 申请金额：50万元 | 期限：24个月 | 有无抵押：有房产(300万)
- 有无逾期：无

## 系统评分
- 综合评分：82/100（低风险）| 企业健康度：85/100
- 五维分项：经营实力18/20 | 现金流覆盖17/20 | 信用合规20/20 | 信用增强15/20 | 杠杆风险12/20

## 银行排名（TOP 3）
1. 邮储银行 通过率 72% 利率 3.05% 额度 240万
2. 平安银行 通过率 68% 利率 3.00% 额度 240万
3. 建设银行 通过率 65% 利率 2.40% 额度 240万

## 相关知识库条目
- 行业准入：制造业 正常准入 接受度系数 1.0，广东省制造业调整系数 +1.15
- 政策：[LOC-RULE-001] 广东省普惠金融实施方案，对单户≤1000万贷款给予风险补偿
- 案例：江苏制造业企业（8年/A级纳税/有房产）获批工商银行经营快贷 80万

请输出：
1. 风险评估（2-3句，说明主要优势和风险点）
2. 银行推荐（说明TOP3银行各自适合的原因，引用具体银行的产品特点）
3. 改善建议（3-5条具体可行的行动建议）
4. 政策提示（结合知识库中的地方政策，告知可享受的优惠）

要求：语言专业但通俗，每条建议都要有数据或政策依据。''')

doc.add_heading('技术选型', level=2)
tbl(doc,
    ['方案', 'API', '成本', '推荐场景'],
    [
        ['DeepSeek V3', 'api.deepseek.com', '极低（1元/百万token）', '首选：中文能力强，价格低'],
        ['Claude (Anthropic)', 'api.anthropic.com', '中', '英文/逻辑推理最强'],
        ['通义千问', 'dashscope.aliyun.com', '低', '备选：阿里云生态'],
        ['本地模型 (Ollama)', 'localhost:11434', '免费', '离线演示：无需网络'],
    ]
)

doc.add_heading('实现要点', level=2)
bullet(doc, '将 API Key 放在 .env 文件中（已在 .gitignore），不要硬编码')
bullet(doc, 'LLM 调用失败时回退到现有的模板诊断，不影响基本功能')
bullet(doc, 'prompt 中嵌入 kb_sources 数据，让 LLM 能引用具体政策编号和案例 ID')
bullet(doc, '在后端做 prompt 构建，前端只负责展示 LLM 返回的 Markdown 文本')

doc.add_page_break()

# ================================================================
# P0-2: 多轮对话
# ================================================================
doc.add_heading('P0-2：多轮对话与 What-if 分析', level=1)

doc.add_heading('现状问题', level=2)
doc.add_paragraph(
    '当前系统是一次性输入 → 一次性输出。用户看完报告后如果有疑问（"为什么邮储比建行高？"'
    '"我把纳税提升到B级会怎样？"），无法追问，只能手动改表单重新提交。'
    '这不符合真实场景——小微企业主在申请贷款前需要反复推演不同方案。'
)

doc.add_heading('改进方案', level=2)
doc.add_paragraph('在前端评估结果页底部增加一个对话栏，支持三种类型的追问：')

bullet(doc, '解释型："为什么 XX 银行通过率更高？" → Agent 分析该银行的产品偏好和用户画像的匹配度')
bullet(doc, 'What-if 型："如果把纳税等级从 C 提升到 B？" → Agent 修改参数重新计算，对比前后差异')
bullet(doc, '建议型："我应该优先准备哪些材料？" → Agent 根据银行要求给出优先级排序')

doc.add_heading('技术实现', level=2)
code(doc, '''# 后端新增 /api/evaluate/chat 端点
class ChatRequest:
    session_id: str          # 会话 ID
    message: str             # 用户追问文本
    context: dict            # 当前评估上下文（评分、银行排名等）

class ChatResponse:
    reply: str               # Agent 回复文本
    updated_context: dict    # 更新后的上下文（what-if 场景下数值会变）
    kb_sources: list         # 本次追问引用的知识库条目''')

doc.add_heading('对话示例', level=2)
code(doc, '''用户：为什么邮储银行通过率比建设银行高？

Agent：主要原因是您的三个特征与邮储更匹配：
1. 您是电商企业（邮储对线上经营数据接受度高，建行偏好传统流水）
2. 您无抵押物（建行要求抵押物严格，邮储纯信用额度 200 万）
3. 您经营2年（建行要求 ≥3 年，邮储仅需 1 年）
[引用 kb: 建行产品要求 min_business_years=3，邮储=1]

用户：那我把纳税等级提升到 A 级呢？

Agent：（重新计算）如果纳税等级从 B 提升到 A：
- 综合评分：60 → 67（+7分）
- 邮储通过率：52% → 61%
- 建行通过率：48% → 55%
- 建设银行银税贷产品对 A 级纳税有专项利率优惠（从 2.40% 降至 2.20%）
[引用 kb: tax_level_scoring.csv A=5, B=4]
建议：如果近期有纳税记录更新，尽快申请纳税等级复评。''')

doc.add_page_break()

# ================================================================
# P1-1: RAG 政策检索
# ================================================================
doc.add_heading('P1-1：RAG 政策检索与引用', level=1)

doc.add_heading('现状问题', level=2)
doc.add_paragraph(
    'kb/data/policies/ 中有 42 条政策（18 条国家级 + 24 条地方级），政策数据极其丰富：'
    '每条包含 rule_id、key_condition、risk_warning、subsidy_content、agent_usage 等字段。'
    '但当前仅在 KBQuery.get_relevant_policies() 中做了简单的字符串包含匹配，检索精度极低。'
)

doc.add_heading('改进方案', level=2)
doc.add_paragraph('三步实现 RAG（检索增强生成）：')

bullet(doc, '1. 向量化：将每条政策的 key_condition + agent_usage + rule_theme 拼接，用 text-embedding 模型生成向量')
bullet(doc, '2. 建索引：存入 ChromaDB（轻量 Python 向量库，零配置）')
bullet(doc, '3. 检索：用户输入（行业 + 地区 + 企业特征）→ 向量化 → 相似度搜索 → 返回 TOP 5 最相关政策')

code(doc, '''# 新增 kb/loader/embedder.py
class PolicyEmbedder:
    def __init__(self):
        self.client = chromadb.PersistentClient(path="./kb/chroma_db")
        self.collection = self.client.get_or_create("policies")

    def build_index(self, kb: KnowledgeBase):
        """从 kb 加载政策并向量化"""
        for _, row in kb.national_policies.iterrows():
            text = f"{row['rule_theme']} {row['key_condition']} {row['agent_usage']}"
            self.collection.add(documents=[text], metadatas=[row.to_dict()], ids=[row['rule_id']])

    def search(self, query: str, province: str = None, top_k=5):
        """语义搜索相关度最高的政策"""
        results = self.collection.query(query_texts=[query], n_results=top_k)
        return results''')

doc.add_heading('效果展示', level=2)
bullet(doc, '银行推荐中引用地方政策，如工行推荐的推荐理由中加入"广东省对单户≤1000万的普惠贷款有风险补偿（粤府办〔2025〕3号）"')
bullet(doc, '改善建议中引用补贴政策，如"作为科创企业，可申请中关村科技型企业融资风险补偿（中科园发〔2025〕7号）"')
bullet(doc, '材料清单中引用政策要求，如"依据流动资金贷款管理办法（金规〔2024〕3号），需提供贷款用途证明材料"')

doc.add_page_break()

# ================================================================
# P1-2: 语义案例匹配
# ================================================================
doc.add_heading('P1-2：语义案例匹配', level=1)

doc.add_heading('现状问题', level=2)
doc.add_paragraph(
    'KBQuery.find_similar_cases() 使用简单的字符串包含匹配（industry + risk_signal + case_type）。'
    '"奶茶店"和"餐饮"可能匹配不上，"流水不足"和"无银行流水"也无法关联。'
)

doc.add_heading('改进方案', level=2)
bullet(doc, '1. 对 20 条增强案例的 summary + risk_signal + industry 字段做语义向量化')
bullet(doc, '2. 用户输入的企业画像做同样处理，计算余弦相似度')
bullet(doc, '3. 返回 TOP 3 语义最接近的案例，附上诊断推理链（diagnosis_chain）和改善建议（improvement_advice）')
bullet(doc, '4. 在诊断报告中展示："和一个您类似的案例——浙江 3 年电商卖家，因为支付宝流水完整，成功获批网商贷 50 万"')

doc.add_heading('为什么重要', level=2)
doc.add_paragraph(
    '对于小微企业主来说，"有个和我差不多的店拿到了贷款"比任何评分数字都有说服力。'
    '案例匹配是这个产品中最有"人味"的功能，也是课程展示时的加分项。'
)

doc.add_page_break()

# ================================================================
# P2-1: PDF 导出
# ================================================================
doc.add_heading('P2-1：评估报告 PDF 导出', level=1)

doc.add_heading('改进方案', level=2)
doc.add_paragraph('在评估结果页增加一个"导出报告"按钮，生成排版精美的 PDF 文件，包含：')
bullet(doc, '封面：企业名称 + 评估日期 + 综合评分')
bullet(doc, '五维评分雷达图（ECharts 渲染 → 截图嵌入）')
bullet(doc, '银行推荐排名表（TOP 10 含通过率、利率、额度、推荐理由）')
bullet(doc, '改善建议清单（优先级排序）')
bullet(doc, '所需材料清单（标注必须 vs 建议）')
bullet(doc, '知识库溯源附录（引用了哪些政策、行业准入、银行产品）')

doc.add_heading('技术方案', level=2)
tbl(doc,
    ['方案', '实现方式', '优点', '缺点'],
    [
        ['前端 jsPDF', 'html2canvas 截图 + jsPDF 拼接', '无需后端改动', '排版控制力弱'],
        ['后端 reportlab', 'Python reportlab 生成', '排版精准，可复现', '需在服务端渲染图表'],
        ['混合方案', '前端收集数据 → 后端用 python-docx 生成 Word', '文档可编辑', '图表需转图片'],
    ]
)
doc.add_paragraph('推荐"混合方案"：Word 比 PDF 更实用，企业主可以拿着去银行、可以编辑补充信息。')

doc.add_page_break()

# ================================================================
# P2-2: 压力测试
# ================================================================
doc.add_heading('P2-2：现金流压力测试', level=1)

doc.add_heading('改进方案', level=2)
doc.add_paragraph(
    '当前系统只评估"当前静态状态"。真实场景中，小微企业最大的风险是现金流波动。'
    '新增 /api/stress-test 端点，模拟三种压力场景：'
)

tbl(doc,
    ['场景', '参数变化', '分析目标'],
    [
        ['营收下降', '月营收 × 0.7, 0.5, 0.3', '月供是否仍能覆盖？断流风险多大？'],
        ['成本上升', '月固定成本 × 1.1, 1.3', '利润被压缩后的还款安全边际'],
        ['回款延迟', '营收周期性波动（旺季/淡季）', '季节性企业的流动性缺口'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('在前端展示为压力测试仪表盘：绿/黄/红三色表示不同压力等级下的还款能力。')

doc.add_page_break()

# ================================================================
# P3: 知识库维护
# ================================================================
doc.add_heading('P3：知识库过期自动提醒', level=1)

doc.add_heading('改进方案', level=2)
doc.add_paragraph('写一个检查脚本 kb/tools/check_staleness.py，运行时可检测：')
bullet(doc, '政策规则的 last_verified_at 超过 90 天未更新')
bullet(doc, '银行产品利率与宏观统计数据中的 LPR 基准偏差过大')
bullet(doc, '数据文件中有空字段或格式错误（基于 JSON Schema）')
bullet(doc, 'kb/VERSION 中的日期与最近一次 git commit 日期不一致')

code(doc, '''# 运行方式
python kb/tools/check_staleness.py

# 示例输出
[OK]  national_policies.csv   18条规则，最近验证：2026-06-20
[WARN] provincial_policies.csv 3条规则超过90天未验证
        LOC-RULE-001 最后验证：2026-03-15
        LOC-RULE-005 最后验证：2026-03-20
[OK]  bank_products.json      28家银行，数据完整
[WARN] kb/VERSION 日期(2026-06-25) 与最近的 git commit(2026-06-25) 不一致''')

doc.add_page_break()

# ================================================================
# Bonus: 多 Agent 协作
# ================================================================
doc.add_heading('Bonus：多 Agent 协作评估', level=1)

doc.add_heading('设计思路', level=2)
doc.add_paragraph(
    '将当前的单一评估流程拆分为 4 个专业子 Agent 并行工作，'
    '最后由汇总 Agent 整合结论。每个子 Agent 只关注一个知识域，'
    '各司其职，互不干扰。这种方式特别适合展示"AI Agent"的课程主题。'
)

doc.add_heading('Agent 分工', level=2)
tbl(doc,
    ['Agent', '职责', '使用的知识库', '输出'],
    [
        ['信用评估 Agent', '分析征信、纳税、经营年限', 'kb/data/credit_tax/', '信用评分 + 风险点'],
        ['银行匹配 Agent', '分析 28 家银行的产品偏好', 'kb/data/banks/', 'TOP 5 银行排名 + 理由'],
        ['政策补贴 Agent', '检索适用的国家和地方政策', 'kb/data/policies/', '可享受的贴息/担保政策'],
        ['案例参考 Agent', '匹配相似的历史案例', 'kb/data/cases/', '3 个相似案例 + 诊断链'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('汇总 Agent 收到四个子 Agent 的报告后，去重、排序、生成最终的诊断报告。')

doc.add_heading('并行优势', level=2)
bullet(doc, '4 个子 Agent 同时运行，总耗时 ≈ 最慢的那个（而非四者之和）')
bullet(doc, '每个子 Agent 的 prompt 更聚焦，幻觉更少')
bullet(doc, '如果某个子 Agent 失败，不影响其他三个——系统有容错')
bullet(doc, '课程报告中更容易展示"多 Agent 协作架构"的技术亮点')

doc.add_page_break()

# ================================================================
# 优先级矩阵
# ================================================================
doc.add_heading('实施优先级矩阵', level=1)

doc.add_paragraph('按"对演示效果的提升"和"实现难度"两个维度排列：')

tbl(doc,
    ['优先级', '功能', '实现难度', '演示冲击力', '建议实施时间'],
    [
        ['🔴 P0 立即做', 'LLM 诊断报告 + 多轮对话', '⭐⭐', '⭐⭐⭐⭐⭐', '本周'],
        ['🟠 P1 尽快做', 'RAG 政策检索', '⭐⭐⭐', '⭐⭐⭐⭐', '下周'],
        ['🟠 P1 尽快做', '语义案例匹配', '⭐⭐', '⭐⭐⭐', '下周'],
        ['🟡 P2 有空做', 'PDF 报告导出', '⭐⭐', '⭐⭐⭐', '课程展示前'],
        ['🟡 P2 有空做', '现金流压力测试', '⭐', '⭐⭐⭐', '课程展示前'],
        ['🟢 P3 日常维护', '知识库过期提醒', '⭐', '⭐⭐', '长期维护'],
        ['🔵 Bonus', '多 Agent 协作', '⭐⭐⭐⭐', '⭐⭐⭐⭐', '课程报告加分项'],
    ]
)

doc.add_heading('推荐的最小可行方案（MVP）', level=2)
doc.add_paragraph(
    '如果时间有限，只做两件事就能实现质的飞跃：'
)
bullet(doc, 'P0-1（LLM 诊断报告）— 把 evalute_loan() 的输出 + kb_query 的检索结果打包成一个 prompt，调 DeepSeek API 生成个性化诊断，替换现有硬编码模板。核心代码约 100 行。')
bullet(doc, 'P0-2（多轮对话）— 在评估结果页底部加一个对话输入框，支持追问，维护对话上下文。核心代码约 150 行。')
doc.add_paragraph('')
doc.add_paragraph(
    '这两个功能加起来大约 2 天工作量，但可以让系统从"高级计算器"变成"真正的 AI Agent"，'
    '课程演示时评委和同学都能直接感知到差异。'
)

doc.add_paragraph()
doc.add_paragraph('—— 文档结束 ——').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---- 保存 ----
output_path = os.path.join(
    os.path.dirname(__file__), '..', '..',
    'Agent功能改进路线图.docx'
)
output_path = os.path.abspath(output_path)
doc.save(output_path)
print(f'Saved: {output_path}')
